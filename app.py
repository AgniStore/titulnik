from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

app = Flask(__name__)

def format_text(text, add_quotes=False):
    # Capitalize first letter
    text = text.strip()
    if text:
        text = text[0].upper() + text[1:]
    # Add quotes if needed
    if add_quotes:
        text = f'"{text}"'
    return text

def split_text_by_length(text, max_length=55):
    if len(text) <= max_length:
        return text, ""
    
    # Find the last space before max_length
    split_index = text[:max_length].rfind(' ')
    if split_index == -1:  # If no space found, force split at max_length
        split_index = max_length
    
    return text[:split_index], text[split_index:].strip()

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_document():
    # Load the template document
    doc = Document('tit.docx')
    
    # Get form data
    form_data = request.form
    
    # Process tables
    for table in doc.tables:
        for i in range(len(table.rows) - 1):  # -1 to avoid last row
            current_row = table.rows[i]
            next_row = table.rows[i + 1]
            
            # Process each cell in the current row
            for j, cell in enumerate(current_row.cells):
                if j < len(next_row.cells):  # Make sure next row has enough cells
                    for paragraph in cell.paragraphs:
                        for key in form_data:
                            if key in paragraph.text:
                                # Format the text
                                formatted_text = format_text(form_data[key], key == "НА ТЕМУ:")
                                
                                if key in ["СПЕЦИАЛЬНОСТЬ", "НА ТЕМУ:"]:
                                    # Split text by length
                                    first_part, second_part = split_text_by_length(formatted_text)
                                    
                                    # Add first part to the right cell
                                    if j + 1 < len(current_row.cells):
                                        # Clear the right cell
                                        for p in current_row.cells[j + 1].paragraphs:
                                            p.clear()
                                        first_paragraph = current_row.cells[j + 1].paragraphs[0]
                                        first_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        run = first_paragraph.add_run(first_part)
                                        run.font.size = Pt(14)
                                    
                                    # Add second part to the cell below
                                    if second_part and j + 1 < len(next_row.cells):
                                        # Clear the cell below
                                        for p in next_row.cells[j + 1].paragraphs:
                                            p.clear()
                                        second_paragraph = next_row.cells[j + 1].paragraphs[0]
                                        second_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        run = second_paragraph.add_run(second_part)
                                        run.font.size = Pt(14)
                                else:
                                    # Insert in the cell below without word wrapping
                                    if j < len(next_row.cells):
                                        # Clear the cell below
                                        for p in next_row.cells[j].paragraphs:
                                            p.clear()
                                        # Add the value to the cell below
                                        paragraph = next_row.cells[j].paragraphs[0]
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                        run = paragraph.add_run(formatted_text)
                                        run.font.size = Pt(14)
    
    # Save the modified document
    output_filename = f'filled_document_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    doc.save(output_filename)
    
    # Send the file to the user
    return send_file(output_filename, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True) 